"""
Create 2-Column Report with Separated Subtasks
==============================================
Clearly separates Subtask 1 and Subtask 2a sections
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print('='*80)
print('Creating 2-Column Separated Report (Subtask 1 | Subtask 2a)')
print('='*80)

# Output path
docx_path = r'D:\Study\Github\Deep-Learning-project-SemEval-2026-Task-2\docs\03_submission\final_submission\PPT, REPORT\SemEval_2026_Task2_Final_Report_Separated.docx'

# Create document
print('\n[1/4] Creating Word document...')
doc = Document()

# Set narrow margins
print('[2/4] Setting page layout (2-column, narrow margins)...')
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Set 2-column layout
def set_column_layout(section, num_columns=2):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), '360')

set_column_layout(sections[0], 2)

# Helper functions
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_subtitle(doc, text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p

def add_section_header(doc, text, color=RGBColor(0, 51, 102)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    return p

def add_text(doc, text, size=8, italic=False, bold=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.font.italic = True
    if bold:
        run.font.bold = True
    return p

def add_bullet(doc, text, size=8):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.15)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def add_compact_table(doc, data, col_widths):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'

    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text)
            if i == 0:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(7)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)

    # Reduce cell padding
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin}')
                node.set(qn('w:w'), '30')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

    return table

# ============================================================================
# CONTENT
# ============================================================================

print('[3/4] Adding content...')

# Title
add_title(doc, 'SemEval 2026 Task 2: Emotional State Change Forecasting')
add_subtitle(doc, 'Final Project Report - Télécom SudParis')
add_subtitle(doc, 'Deep Learning Project | Submission: January 28, 2026', size=8)
add_text(doc, 'Rostislav SVITSOV (Subtask 1: Between-User) • Hyun Chang-Yong 현창용 (Subtask 2a: Within-User)', size=8, italic=True, space_after=6)

# ============================================================================
# 1. DATA PREPARATION
# ============================================================================
add_section_header(doc, '1. DATA PREPARATION')

add_subsection(doc, 'Dataset Overview (Shared)')
add_text(doc, 'SemEval 2026 Task 2: longitudinal ecological essays with self-reported emotions (2021-2024). Dimensions: Valence (pleasantness, 0-4), Arousal (energy, 0-4).')

# SUBTASK 1
add_subsection(doc, '📍 Subtask 1: Between-User Data')
add_text(doc, '[Rostislav will complete: cross-user data characteristics, preprocessing for between-user prediction]', italic=True, size=7.5)

# SUBTASK 2A
add_subsection(doc, '📍 Subtask 2a: Within-User Data')
add_bullet(doc, 'Train: 137 users | Test: 46 users (→1,266 predictions)')
add_bullet(doc, '~25 entries/user over 3 years, irregular time intervals')
add_bullet(doc, 'Valence: μ=2.1,σ=0.9 (balanced) | Arousal: μ=1.0,σ=0.6 (low var)')
add_bullet(doc, 'Critical finding: Arousal 38% less frequent but more sporadic')

add_subsection(doc, '📍 Subtask 2a: Preprocessing')
add_bullet(doc, 'RoBERTa tokenization (max 128 tokens), multilingual')
add_bullet(doc, 'Temporal: lag(t-1,2,3), rolling stats(μ,σ w=3,5,10), log(time gaps)')
add_bullet(doc, 'Arousal-specific: Δ, volatility, acceleration (3-dim)')
add_bullet(doc, 'User embeddings: 64-dim learnable per user')
add_bullet(doc, 'Missing: forward-fill + temporal interpolation')

# ============================================================================
# 2. METHODOLOGY
# ============================================================================
add_section_header(doc, '2. METHODOLOGY')

# SUBTASK 1
add_subsection(doc, '📍 Subtask 1: Between-User Approach')
add_text(doc, '[Rostislav: model architecture, training config, key techniques for cross-user generalization]', italic=True, size=7.5)

# SUBTASK 2A
add_subsection(doc, '📍 Subtask 2a: Hybrid Architecture')
add_bullet(doc, '①RoBERTa-base (125M): 768-dim CLS embeddings')
add_bullet(doc, '②BiLSTM: 2×256 bidir (→512-dim), temporal modeling')
add_bullet(doc, '③Multi-head Attn: 8 heads×96-dim, emotional salience')
add_bullet(doc, '④Features: LSTM(512)+User(64)+Temporal(20)+Text(15)=611-dim')
add_bullet(doc, '⑤Dual-head: Separate Valence/Arousal [611→256→128→1]')

add_subsection(doc, '📍 Subtask 2a: Training')
add_bullet(doc, 'AdamW lr=1e-5, warmup+cosine, batch=16, epochs=50, dropout=0.3')
add_bullet(doc, 'Loss: Weighted CCC+MSE. Standard: 65%+35%(V), 70%+30%(A)')
add_bullet(doc, 'Arousal Specialist: 50%+50%(V), 90%CCC+10%MSE(A)⭐')
add_bullet(doc, 'Colab Pro A100 40GB, FP16 mixed precision')

add_subsection(doc, '⭐ Subtask 2a Innovation: Arousal-Specialist')
add_text(doc, 'Problem: 38% CCC gap (V vs A). Causes: (1)subjective variance—energy harder to report than mood; (2)low variation—loss ignores subtle arousal.', size=7.5)
add_text(doc, 'Solution: ①90% CCC loss (forces agreement), ②dimension-specific params, ③weighted sampling (oversample high-Δ events).', size=7.5, bold=True)
add_text(doc, 'Result: Arousal CCC 0.5281→0.5832 (+6.0%)', size=7.5, bold=True)

add_subsection(doc, '📍 Subtask 2a: Ensemble')
add_bullet(doc, '5 models (seeds 42,123,777,888,1111-arousal), 5000+ combos')
add_bullet(doc, 'Final 2-model (50:50): seed777(V master)+arousal(A master)')
add_bullet(doc, '2-model > 3-5 models (quality beats quantity)')

# ============================================================================
# 3. RESULTS
# ============================================================================
add_section_header(doc, '3. RESULTS AND CONCLUSION')

# SUBTASK 1
add_subsection(doc, '📍 Subtask 1: Results')
add_text(doc, '[Rostislav: performance metrics, key findings]', italic=True, size=7.5)

# SUBTASK 2A
add_subsection(doc, '📍 Subtask 2a: Performance')
add_text(doc, '✅ CCC: 0.6833 (Target 0.62, +10.4%) | Valence 0.7834 | Arousal 0.5832', bold=True, size=8)

perf_table = [
    ['Phase', 'Config', 'CCC', 'Δ', 'Date'],
    ['1', '3-model', '0.6046', 'Base', 'Nov'],
    ['2', '2-model', '0.6305', '+4.3%', 'Dec'],
    ['3', '777+888', '0.6687', '+10.6%', 'Dec23'],
    ['4', '777+aro⭐', '0.6833', '+13%', 'Dec24'],
]
add_compact_table(doc, perf_table, [0.4, 1.0, 0.5, 0.4, 0.6])

add_text(doc, ' ', size=3)

bench_table = [
    ['Model', 'CCC', 'Val', 'Aro', 'Time'],
    ['seed777', '.6554', '.7593', '.5516', '2.5h'],
    ['arousal⭐', '.6512', '.7192', '.5832', '24m'],
    ['seed888', '.6211', '.7210', '.5212', '2.3h'],
]
add_compact_table(doc, bench_table, [0.8, 0.5, 0.5, 0.5, 0.5])

add_subsection(doc, '📍 Subtask 2a: Key Findings')
add_bullet(doc, 'Dimension-specific optimization: +4.3% vs multi-tasking')
add_bullet(doc, '90% CCC weighting critical for agreement metrics')
add_bullet(doc, '2-model > 3-5 models (quality over quantity validated)')
add_bullet(doc, 'Temporal features: +6% accuracy | Multi-seed: -12% variance')

add_subsection(doc, 'Test Submission')
add_text(doc, 'Subtask 2a: 46 users, 1,266 predictions, expected test CCC: 0.67-0.69.', size=7.5)

# ============================================================================
# 4. INDIVIDUAL CONTRIBUTIONS
# ============================================================================
add_section_header(doc, '4. INDIVIDUAL CONTRIBUTIONS')

# SUBTASK 1
add_subsection(doc, '📍 Rostislav SVITSOV - Subtask 1')
add_text(doc, 'Task: Between-User Emotion Recognition (predicting emotions for different users using shared linguistic cues).', italic=True, size=7.5)
add_text(doc, '[To be completed by team member: data prep, model design, training, evaluation, results, time investment, skills]', italic=True, size=7)

# SUBTASK 2A
add_subsection(doc, '📍 Hyun Chang-Yong 현창용 - Subtask 2a (100%)')

add_text(doc, '①Data Analysis (3 months): EDA→38% Arousal gap, 47-dim features (temporal/textual/personal), arousal-specific Δ/volatility/accel.', size=7.5)

add_text(doc, '②Architecture Design (2 months): RoBERTa+BiLSTM+Attn hybrid, dual-head output, Arousal-Specialist (90% CCC loss), 5 models trained.', size=7.5)

add_text(doc, '③Ensemble Optimization (1 month): 5,000+ weight combos, optimal 2-model (50:50), quality>quantity validated.', size=7.5)

add_text(doc, '④Production Pipeline (1 month): Colab 9-step notebook, dynamic dims(863↔866), 1,266 test predictions, submission.zip.', size=7.5)

add_text(doc, '⑤Documentation (ongoing): 40-pg report (FINAL_REPORT.md), 6 guides (train/predict/eval), project reorganization (01/02/03 phases).', size=7.5)

add_text(doc, '⭐Innovation: Arousal-Specialist solved 38% gap via 90% CCC loss engineering → CCC 0.6833 (+10.4% target). Proved dim-specific beats multi-task by 4.3%.', bold=True, size=7.5)

add_text(doc, 'Time: 15 months (Nov 24-Jan 26) | GPU: 10h (A100) | Code: 721 lines | Docs: 2,500+ lines', size=7, italic=True)

add_text(doc, 'Skills: PyTorch, Transformers, RoBERTa, BiLSTM, ensemble, grid search, FP16, Git', size=7, italic=True)

# ============================================================================
# CONCLUSION
# ============================================================================
add_section_header(doc, 'CONCLUSION')
add_text(doc, 'Subtask 2a: CCC 0.6833 (+10.4% vs 0.62). Arousal-Specialist (90% CCC) solved 38% gap, proving dimension-specific optimization beats multi-tasking. 2-model ensemble: quality > quantity. Production pipeline submitted to SemEval 2026.', size=8)

# Footer
add_text(doc, ' ', size=3)
add_text(doc, 'Code: github.com/ThickHedgehog/Deep-Learning-project-SemEval-2026-Task-2 | Competition: codabench.org/competitions/9963/', size=6, italic=True, space_after=0)

# Save
print('[4/4] Saving document...')
doc.save(docx_path)

print('\n' + '='*80)
print('✅ SUCCESS!')
print('='*80)
print(f'\n📄 Separated Report: {docx_path}')
print(f'\n📊 Structure:')
print(f'   - Format: 2-column, ~2 pages')
print(f'   - Subtask 1: Clearly marked with 📍 (Rostislav to complete)')
print(f'   - Subtask 2a: Fully detailed (Hyun Chang-Yong)')
print(f'   - Sections: Data prep, Methodology, Results, Contributions')
print('\n💡 Separation strategy:')
print('   📍 Subtask 1 sections: Placeholders for team member')
print('   📍 Subtask 2a sections: Complete with all details')
print('   ✅ Clear visual distinction (📍 markers)')
print('\n' + '='*80)
